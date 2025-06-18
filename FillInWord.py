from docx import Document
import re

def fill_in_word(doc_name, output_name, answers):
    """
    Takes in a word doc and an array of answers, filling
    in the responses into the doc. Right now, fills the answers
    into the [Answer X] boxes
    """
    doc = Document(doc_name)

    index = 0

    # cycle through each paragraph, replacing if [Answer X] is in it
    for paragraph in doc.paragraphs:
        placeholder = f'[Answer {index + 1}]'

        if(placeholder in paragraph.text):
            paragraph.text = paragraph.text.replace(placeholder, answers[index])
            index += 1

    # once all finished, save
    doc.save(output_name)


def grab_sections(doc_name):
    """
    Take in a word doc and grabs the sections preceding each answer.
    These sections can be used within prompt to enhance the retrieval.
    """
    doc = Document(doc_name)

    sections = []
    current_section = []

    # matches [Answer X] sections
    answer_pattern = re.compile(r"\[Answer\s*\d+\]", re.IGNORECASE)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        # if empty paragraph, continue
        if(not text):
            continue

        if(answer_pattern.match(text)):
            # if answer, add current section to list of actual sections
            if(current_section):
                section_text = " ".join(current_section).strip()
                sections.append(section_text)
                current_section = []
        else:
            # if not answer, add text to current section
            current_section.append(text)

    return sections


def grab_n_shot(doc_name, num_examples=3):
    """
    Grabs n_shot examples from a word document in the form of:

    [Section N]
    This is a really cool section!
    [Response N]
    Woah this is an AMAZING response!

    Adding examples of similar section->response pairs can drastically increase performance.
    If there is no word doc in this format, the examples can be added directly into the system
    prompt to simulate the same process.
    If you want to see specific formats (such as the use of bullet points, or bold text) you
    MUST include those in the n_shot, otherwise the system prompt must be configured properly.
    The more relevant your n_shot examples to the area you are autofilling, the better performance
    you can expect from this bot.

    """
    doc = Document(doc_name)

    n_shot_examples = []
    current_section = []

    # if this is true, matches a section. When reaching a response, turns false, until reaching another section
    match_section = True

    # matches [Section X] sections
    section_pattern = re.compile(r"\[Section\s*\d+\]", re.IGNORECASE)
    # matches [Response X] sections
    response_pattern = re.compile(r"\[Response\s*\d+\]", re.IGNORECASE)


    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        # if empty paragraph, continue
        if(not text):
            continue

        if(section_pattern.match(text) and match_section):
            # this triggers when starting on the first section
            current_section.append("Section:\n")
        elif(section_pattern.match(text) and not match_section):
            # this triggers on each suqsequent section
            n_shot_examples.append(" ".join(current_section).strip())
            current_section = []
            current_section.append("Section:\n")
            match_section = True
        elif(response_pattern.match(text)):
            # this triggers on response
            match_section = False
            current_section.append("\nResponse:\n")
        else:
            # if no pattern matches, it must be a normal paragraph
            current_section.append(text)

    # Aaaaannnnnndd don't forget to add the last example at the end:
    n_shot_examples.append(" ".join(current_section).strip())
    # The fact that I have to do the above is probably a bad code alert haha

    examples = "\n".join(n_shot_examples[:num_examples]).strip()
    return examples



if __name__ == "__main__":
    print(grab_sections(r"C:\Users\DSU\Research\GPTDOCs\AutoFillBot\AppleDoc.docx"))
    print(grab_n_shot(r"C:\Users\DSU\Research\GPTDOCs\AutoFillBot\N_shot_examples.docx", 3))
